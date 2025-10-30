# app.py — Data Profiling AI Assistant (v4.0, Multi-File Support)
import streamlit as st
import pandas as pd
import json
import time
import io
import pdfplumber
from docx import Document
import xml.etree.ElementTree as ET
from datetime import datetime

# Import custom modules
from profiler import profile_dataframe
from pbl_generator import generate_pbl_for_column, derive_rules_from_reference
from chat_agent import answer_question_about_df

# ----------------------------
# Branding
APP_TITLE = "Data Profiling AI Assistant 🤖"
APP_SUBTITLE = "Smart insights, business logic & AI-powered data understanding"
PRIMARY_COLOR = "#007BFF"  # Corporate blue
BACKGROUND_GRADIENT = "linear-gradient(135deg, #E3F2FD, #BBDEFB)"  # Light blue gradient
# ----------------------------

st.set_page_config(page_title=APP_TITLE, page_icon="🤖", layout="wide")

# ----------------------------
# CSS Styling
st.markdown(f"""
<style>
body {{
    background: {BACKGROUND_GRADIENT};
    background-attachment: fixed;
    color: #0D1B2A;
    font-family: 'Segoe UI', sans-serif;
}}
.app-header {{
    text-align: center;
    padding: 20px 0 10px 0;
}}
.app-title {{
    color: {PRIMARY_COLOR};
    font-size: 32px;
    font-weight: 800;
    margin-bottom: 4px;
}}
.app-subtitle {{
    color: #0D1B2A;
    font-size: 15px;
    font-weight: 400;
}}
.card {{
    background: rgba(255, 255, 255, 0.9);
    border: 1px solid rgba(0, 123, 255, 0.2);
    border-radius: 16px;
    padding: 18px;
    margin-bottom: 18px;
    box-shadow: 0 4px 12px rgba(0,0,0,0.08);
}}
h3, h4 {{
    color: {PRIMARY_COLOR};
    margin-top: 10px;
}}
.stProgress > div > div > div > div {{
    background-color: {PRIMARY_COLOR};
}}
</style>
""", unsafe_allow_html=True)

# ----------------------------
# Header
st.markdown("<div class='app-header'>", unsafe_allow_html=True)
st.markdown(f"""
<div class='app-title'>{APP_TITLE}</div>
<div class='app-subtitle'>{APP_SUBTITLE}</div>
<div style='margin-top:10px;'>
    <span style='background-color:#007BFF;
                 color:white;
                 padding:6px 14px;
                 border-radius:12px;
                 font-size:13px;
                 font-weight:600;'>
    🤖 Powered by Gemini-2.5 flash · Generative AI Layer
    </span>
</div>
""", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)
st.markdown("---")

# ----------------------------
# Sidebar Uploaders
with st.sidebar:
    st.header("📂 Upload Datasets")
    uploaded_file = st.file_uploader(
        "Upload Target File (CSV, XLSX, JSON, TXT, PDF, DOCX, XML)",
        type=["csv", "xlsx", "xls", "json", "txt", "pdf", "docx", "xml"],
        key="u_target"
    )

    ref_file = st.file_uploader(
        "Upload Reference File (Optional)",
        type=["csv", "xlsx", "xls", "json", "txt", "pdf", "docx", "xml"],
        key="u_ref"
    )

    st.markdown("---")
    st.caption("💡 Use small files (<5MB) for best performance.")
    st.caption("Built for Hackathon 2025 | Streamlit")

# ----------------------------
# Multi-format file loader
def load_dataframe(uploaded):
    """Reads CSV, Excel, JSON, TXT, PDF, DOCX, or XML into a pandas DataFrame."""
    if uploaded is None:
        return None

    file_name = uploaded.name.lower()
    try:
        # CSV
        if file_name.endswith(".csv"):
            return pd.read_csv(uploaded)

        # Excel
        elif file_name.endswith((".xls", ".xlsx")):
            return pd.read_excel(uploaded)

        # JSON
        elif file_name.endswith(".json"):
            data = json.load(uploaded)
            if isinstance(data, list):
                return pd.DataFrame(data)
            elif isinstance(data, dict):
                return pd.json_normalize(data)
            else:
                st.warning("Unsupported JSON structure.")
                return None

        # TXT
        elif file_name.endswith(".txt"):
            content = uploaded.read().decode("utf-8", errors="ignore")
            if "\t" in content:
                return pd.read_csv(io.StringIO(content), sep="\t")
            elif "|" in content:
                return pd.read_csv(io.StringIO(content), sep="|")
            else:
                return pd.read_csv(io.StringIO(content))

        # PDF (extract tables)
        elif file_name.endswith(".pdf"):
            with pdfplumber.open(uploaded) as pdf:
                tables = []
                for page in pdf.pages:
                    table = page.extract_table()
                    if table:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables.append(df)
                if tables:
                    return pd.concat(tables, ignore_index=True)
                else:
                    st.warning("No tables detected in the PDF file.")
                    return None

        # Word Document (.docx)
        elif file_name.endswith(".docx"):
            doc = Document(uploaded)
            tables = []
            for table in doc.tables:
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0])
                tables.append(df)
            if tables:
                return pd.concat(tables, ignore_index=True)
            else:
                text = "\n".join([p.text for p in doc.paragraphs])
                st.text_area("Extracted Text from Word File", text[:5000])
                return None

        # XML
        elif file_name.endswith(".xml"):
            tree = ET.parse(uploaded)
            root = tree.getroot()
            rows, columns = [], set()
            for elem in root:
                row = {child.tag: child.text for child in elem}
                rows.append(row)
                columns.update(row.keys())
            return pd.DataFrame(rows, columns=sorted(columns))

        else:
            st.error("Unsupported file type.")
            return None

    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

df = load_dataframe(uploaded_file) if uploaded_file else None
ref_df = load_dataframe(ref_file) if ref_file else None

# ----------------------------
# Tabs
tabs = st.tabs(["📊 Overview", "⚙️ Profiling & PBL", "💬 Chat Assistant", "📥 Export"])

# ----------------------------
# Tab 1: Overview
with tabs[0]:
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("Dataset Overview")
    if df is None:
        st.info("👆 Upload a dataset using the sidebar to begin.")
    else:
        st.success(f"Loaded dataset with {df.shape[0]} rows and {df.shape[1]} columns.")
        st.dataframe(df.head(200))
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Rows", f"{len(df):,}")
        col2.metric("Columns", len(df.columns))
        col3.metric("Nulls", int(df.isna().sum().sum()))
        col4.metric("Distinct Values", int(sum(df.nunique())))
    st.markdown("</div>", unsafe_allow_html=True)

# ----------------------------
# Tab 2: Profiling & PBL
with tabs[1]:
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("Profiling & Business Logic Generation")
    if df is None:
        st.info("Upload a target dataset first.")
    else:
        progress_text = "Profiling dataset... please wait ⏳"
        progress_bar = st.progress(0, text=progress_text)
        for pct in range(0, 100, 10):
            time.sleep(0.1)
            progress_bar.progress(pct + 10, text=progress_text)
        with st.spinner("Analyzing columns and computing statistics..."):
            profile = profile_dataframe(df)
        progress_bar.empty()
        st.success("✅ Profiling completed successfully!")

        st.subheader("📈 Summary Statistics")
        overview_rows = []
        for c, meta in profile["columns_overview"].items():
            overview_rows.append({
                "Column": c,
                "Type": meta.get("dtype"),
                "Nulls": meta.get("nulls"),
                "Uniqueness": meta.get("distinct"),
                "Pattern": meta.get("pattern"),
                "Min": meta.get("min"),
                "Max": meta.get("max"),
                "Outliers": meta.get("outliers", {}).get("count", 0)
            })
        st.dataframe(pd.DataFrame(overview_rows))
        
        # Null and unique counts
        st.subheader("🧮 Null & Unique Counts")
        nulls = df.isnull().sum().sort_values(ascending=False)
        uniques = df.nunique().sort_values(ascending=False)

        col1, col2 = st.columns(2)
        with col1:
            st.bar_chart(nulls)
        with col2:
            st.bar_chart(uniques)

        st.markdown("### Generate PBL (Auto)")
        col = st.selectbox("Select column for PBL", df.columns.tolist())
        if st.button("Generate PBL"):
            with st.spinner("Creating business logic..."):
                rules = generate_pbl_for_column(df[col], col_name=col)
                time.sleep(1)
            st.success(f"PBL for `{col}`:")
            for i, r in enumerate(rules, 1):
                st.write(f"{i}. {r}")

        # Reference File Comparison
        if ref_df is not None:
            st.markdown("---")
            st.markdown("### Derive PBL from Reference File")
            st.caption("Compare and generate PBL based on reference dataset mappings.")
            ref_cols = ref_df.columns.tolist()
            mapping = {}
            for c in df.columns:
                default = c if c in ref_cols else None
                mapping[c] = st.selectbox(f"Map `{c}` to:", [None] + ref_cols,
                                          index=(ref_cols.index(default)+1 if default else 0), key=f"map_{c}")

            if st.button("Derive Rules from Reference"):
                with st.spinner("Analyzing reference vs target..."):
                    progress_bar = st.progress(0, text="Deriving rules...")
                    results = {}
                    for i, (tgt, ref) in enumerate(mapping.items()):
                        time.sleep(0.1)
                        if ref:
                            results[tgt] = derive_rules_from_reference(df[tgt], ref_df[ref])
                        progress_bar.progress(int(((i+1)/len(mapping))*100))
                    progress_bar.empty()
                st.success("✅ Derived rules successfully.")
                for k, rules in results.items():
                    st.markdown(f"**{k}**")
                    for r in rules:
                        st.write("-", r)
    st.markdown("</div>", unsafe_allow_html=True)

# ----------------------------
# Tab 3: Chat Assistant
with tabs[2]:
    #st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("Chat with Assistant 🤖")
    if df is None:
        st.info("Upload a dataset first.")
    else:
        q = st.text_input("Ask about the data (e.g., 'Which columns have nulls?')")
        if st.button("Ask"):
            with st.spinner("Thinking..."):
                ans = answer_question_about_df(q, df, profile)
                time.sleep(0.8)
            st.success("Answer:")
            st.write(ans)
    st.markdown("</div>", unsafe_allow_html=True)

# ----------------------------
# Tab 4: Export
with tabs[3]:
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("Export Results 📥")
    if df is None:
        st.info("Upload a dataset first.")
    else:
        profile_json = json.dumps(profile["columns_overview"], indent=2, default=str)
        st.download_button("⬇️ Download Profiling JSON", data=profile_json, file_name="profiling_summary.json")
    st.markdown("</div>", unsafe_allow_html=True)

# ----------------------------
# Floating Chatbot Icon
chat_icon_html = f"""
<style>
.chat-float {{
  position: fixed;
  bottom: 25px;
  right: 25px;
  z-index: 9999;
}}
.chat-float button {{
  background-color: {PRIMARY_COLOR};
  border: none;
  color: white;
  font-size: 26px;
  width: 60px;
  height: 60px;
  border-radius: 50%;
  box-shadow: 0 4px 8px rgba(0,0,0,0.3);
  cursor: pointer;
  animation: pulse 2.5s infinite;
}}
@keyframes pulse {{
  0% {{ box-shadow: 0 0 0 0 rgba(0,123,255,0.5); }}
  70% {{ box-shadow: 0 0 0 12px rgba(0,123,255,0); }}
  100% {{ box-shadow: 0 0 0 0 rgba(0,123,255,0); }}
}}
.chat-float button:hover {{
  transform: scale(1.1);
  box-shadow: 0 4px 12px rgba(0,123,255,0.6);
}}
</style>

<div class="chat-float">
  <button title="Chat Assistant 💬" onClick="window.scrollTo(0, document.body.scrollHeight);">💬</button>
</div>
"""
st.markdown(chat_icon_html, unsafe_allow_html=True)

# ----------------------------
st.markdown("<hr>", unsafe_allow_html=True)
st.caption("💼 Built for Hackathon 2025 | Streamlit | v4.0 (Multi-Format + Blue Theme)")
